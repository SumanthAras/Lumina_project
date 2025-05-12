from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, send_file, abort
from flask_cors import CORS
from flask_mail import Mail, Message
from dotenv import load_dotenv
from pymongo import MongoClient
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from itsdangerous import URLSafeTimedSerializer as Serializer
import os
import pandas as pd
import numpy as np
from scipy import stats
import statsmodels.api as sm
import matplotlib.pyplot as plt
from sklearn.decomposition import FactorAnalysis
from io import BytesIO
import gridfs
from bson import ObjectId
from bson.errors import InvalidId
from scipy.stats import chi2_contingency
from flask_wtf import FlaskForm
from wtforms import StringField, PasswordField
from wtforms.validators import InputRequired, Email
from itsdangerous import URLSafeTimedSerializer
from authlib.integrations.flask_client import OAuth
import smtplib, ssl
from bson.objectid import ObjectId
import requests
from flask import send_from_directory
from email.message import EmailMessage
from markupsafe import Markup
from docx import Document
from datetime import datetime
from passlib.hash import scrypt

# ✅ Load environment variables from .env file
load_dotenv()

# ✅ Initialize the Flask app
app = Flask(__name__)
CORS(app)

# ✅ Secret key for sessions
app.secret_key = os.getenv("FLASK_SECRET_KEY", "your_secret_key")

# ✅ Mail configuration using environment variables
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = 'luminaassitant@gmail.com'
app.config['MAIL_PASSWORD'] = 'hfpe opim alde izqh'
app.config['MAIL_DEFAULT_SENDER'] = 'luminaassitant@gmail.com'

# ✅ Initialize Flask-Mail
mail = Mail(app)


# Load environment variables
SENDER_EMAIL = os.getenv("SENDER_EMAIL")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD")
CLIENT_SECRETS_FILE = "client_secret.json"
SCOPES = ["https://www.googleapis.com/auth/forms.body"]
MONGO_URI = os.getenv("MONGO_URI")# Make sure to add this to your .env file#

from urllib.parse import quote_plus

username = "LuminaAssistant"
password = quote_plus("Lumina@001")

client = MongoClient("mongodb+srv://LuminaAssistant:Lumina%40001@cluster0.aaqfige.mongodb.net/?retryWrites=true&w=majority&appName=Cluster0")
db = client["lumina_db"]  # Database Name
users_collection = db["users"]  # Collection Name
fs = gridfs.GridFS(db)
pdfs_collection = db["pdfs"]
admin_collection = db["admin_users"]
validators_collection = db["validators"]
pending_users_collection = db["pending_users"]
query_collection = db["queries"]
feedback_collection = db["feedbacks"]
assessment_collection = db["assessments"]



# File Upload Config
UPLOAD_FOLDER = 'static/uploads/'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

ALLOWED_EXTENSIONS = {'jpg', 'jpeg', 'png', 'pdf'}
# Global storage
result_text = ""
plot_stream_buffer = None

def alert_and_redirect(message, route):
    script = f"""
    <script>
        alert("{message}");
        window.location.href = "{url_for(route)}";
    </script>
    """
    return Markup(script)

oauth = OAuth(app)

google = oauth.register(
    name='google',
    client_id='1099459672351-al23dc9fr9noffeuqmhrc4cgcs5el4ph.apps.googleusercontent.com',
    client_secret='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',
    access_token_url='https://accounts.google.com/o/oauth2/token',
    access_token_params=None,
    authorize_url='https://accounts.google.com/o/oauth2/auth',
    authorize_params=None,
    api_base_url='https://www.googleapis.com/oauth2/v1/',
    userinfo_endpoint='https://www.googleapis.com/oauth2/v1/userinfo',  # Needed for fetching user info
    client_kwargs={'scope': 'openid email profile'},
)

@app.route('/')
def home():
    if "email" in session:
        user = users_collection.find_one({"email": session["email"]})
        if user and user.get("image_id"):
            profile_img = url_for('get_profile_pic', image_id=user["image_id"])
        else:
            profile_img = get_default_profile_image()
    else:
        # Not logged in: show default image and possibly redirect or show login option
        user = None
        profile_img = get_default_profile_image()
    return render_template('index.html', profile_img=profile_img, user=user)

@app.route('/admin-login', methods=['GET', 'POST'])  # <-- must include POST here
def admin_login():
    if request.method == 'POST':
        admin_email = request.form['admin_email']
        admin_password = request.form['admin_password']

        admin = admin_collection.find_one({'admin_email': admin_email})
        if admin and check_password_hash(admin['admin_password'], admin_password):
            session['admin_email'] = admin_email
            session['admin_name'] = admin.get('admin_name', 'Admin')
            return redirect('/admin-dashboard')
        else:
            return render_template('validator_login.html', alert="Invalid password or email")

    return render_template('admin_login.html')

@app.route('/admin-dashboard')
def admin_dashboard():
    if 'admin_email' not in session:
        return redirect('/admin-login')

    admin_name = session.get('admin_name', 'Admin')

    # Count total users and validators
    total_users = users_collection.count_documents({})
    total_validators = validators_collection.count_documents({})
    total_admins= admin_collection.count_documents({})
    queries = query_collection.count_documents({})
    feedback = feedback_collection.count_documents({})
    pending_users_count = db["pending_users"].count_documents({})

    pending_users = list(db["pending_users"].find({}))

    return render_template(
        'admin_dashboard.html',
        admin_name=admin_name,
        total_users=total_users,
        pending_users_count=pending_users_count,
        total_validators=total_validators,
        total_admins=total_admins,
        queries=queries,
        feedback = feedback,
        pending_users=pending_users
    )


@app.route('/validator-dashboard')
def validator_dashboard():
    validator_email = session.get('validator_email')  # corrected
    name = session.get('validator_name')  # corrected

    if not validator_email:
        return "Unauthorized", 403

    validator = validators_collection.find_one({'email': validator_email})

    if not validator:
        return "Unauthorized", 403

    total_validators = validators_collection.count_documents({})
    total_admins = admin_collection.count_documents({})
    assessments = assessment_collection.count_documents({})

    categories = validator.get('categories', [])

    entries = list(assessment_collection.find({
        'category': {'$in': categories}
    }))

    return render_template(
        'validator_dashboard.html',
        name=name,
        entries=entries,
        total_validators=total_validators,
        total_admins=total_admins,
        assessments=assessments
    )

@app.route('/uploads/<filename>')
def download_file(filename):
    return send_from_directory(UPLOAD_FOLDER, filename, as_attachment=True)

@app.route('/validate_save', methods=['POST'])
def validate_save():
    data = request.get_json()
    email = data.get("email")
    category = data.get("category")
    questionnaire = data.get("questionnaire", [])

    if not email or not category or not questionnaire:
        return jsonify({"error": "Missing required data"}), 400

    # Generate Word doc
    doc = Document()
    doc.add_heading(f"{category} Questionnaire", level=1)

    for idx, item in enumerate(questionnaire, 1):
        doc.add_paragraph(f"{idx}. {item['question']}", style='List Number')
        for opt in item['options']:
            doc.add_paragraph(f"- {opt}", style='List Bullet')

    filename = f"{email.replace('@', '')}{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    doc.save(filepath)

    # Save to MongoDB
    assessment_collection.insert_one({
        "email": email,
        "category": category,
        "filename": filename,
        "created_at": datetime.utcnow()
    })

    return jsonify({"message": "Success", "filename": filename})

@app.route('/submit-evaluated', methods=['POST'])
def submit_evaluated():
    file = request.files.get('evaluated_file')
    email = request.form.get('email')
    category = request.form.get('category')

    if not file or not email or not category:
        return jsonify({"error": "Missing data"}), 400

    filename = f"evaluated_{email.replace('@', '')}{datetime.now().strftime('%Y%m%d%H%M%S')}.pdf"
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # Optional: Store or update record in MongoDB
    assessment_collection.update_one(
        {"email": email, "category": category},
        {"$set": {"evaluated_file": filename, "evaluated_at": datetime.utcnow()}}
    )

    # Send email with evaluated file (mock implementation for now)
    print(f"Would send {filename} to {email}")

    return jsonify({"message": "Evaluated file received and stored."})


@app.route('/api/validator-dashboard/get-validators')
def get_validators_for_validator():
    if 'validator_email' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    validators = list(validators_collection.find({}, {'_id': 0, 'name': 1, 'email': 1, 'categories': 1}))
    return jsonify(validators)


# Route to get list of admins (for validator dashboard)
@app.route('/api/validator-dashboard/get-admins')
def get_admins_for_validator():
    if 'validator_email' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    admins = list(admin_collection.find({}, {'_id': 0, 'admin_name': 1, 'admin_email': 1}))
    return jsonify(admins)

@app.route('/validator-login', methods=['GET', 'POST'])
def validator_login():
    if request.method == 'POST':
        form_type = request.form.get('form_type')

        if form_type == 'register':
            name = request.form['validator_name']
            email = request.form['validator_email']
            password = request.form['validator_password']
            categories = request.form.getlist('categories[]')

            existing_validator = validators_collection.find_one({'email': email})

            if existing_validator:
                return render_template('validator_login.html', alert="Account already exists. Please login.")

            hashed_password = generate_password_hash(password)
            validators_collection.insert_one({
                'name': name,
                'email': email,
                'password': hashed_password,
                'categories': categories
            })

            return render_template('validator_login.html', alert="Registration successful! Please log in.")

        elif form_type == 'login':
            email = request.form['validator_email']
            password = request.form['validator_password']

            validator = validators_collection.find_one({'email': email})

            if not validator:
                return render_template('validator_login.html', alert="Kindly register first!")

            if check_password_hash(validator['password'], password):
                session['validator_email'] = validator['email']
                session['validator_name'] = validator['name']
                return redirect('/validator-dashboard')
            else:

                return render_template('validator_login.html', alert="Incorrect email or password!")

    return render_template('validator_login.html')

def get_pending_users():
    return list(users_collection.find({"status": "pending"}))

@app.route('/manage-users')
def manage_users():
    pending_users = get_pending_users()
    return render_template('manage_users.html', pending_users=pending_users)


@app.route('/approve_user', methods=['POST'])
def approve_user():
    data = request.get_json()  # Get the JSON data sent from JavaScript
    user_id = data.get('user_id')
    user = pending_users_collection.find_one({"_id": ObjectId(user_id)})

    if user:
        users_collection.insert_one({
            "name": user["name"],
            "email": user["email"],
            "password": user["password"],
            "role": user.get("role", "user"),
            "status": "approved"
        })
        pending_users_collection.delete_one({"_id": ObjectId(user_id)})
        return jsonify({"message": "✅ User approved successfully"}), 200

    return jsonify({"message": "User not found"}), 404


@app.route('/reject_user', methods=['POST'])
def reject_user():
    data = request.get_json()
    user_id = data.get("user_id")
    result = pending_users_collection.update_one(
        {"_id": ObjectId(user_id)},
        {"$set": {"status": "rejected"}}
    )
    pending_users_collection.delete_one({"_id": ObjectId(user_id)})
    return jsonify({"message": "❌ User rejected"}), 200

@app.route('/api/get-users')
def get_users():
    if 'admin_email' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    users = list(users_collection.find({}, {'_id': 0, 'name': 1, 'email': 1}))
    return jsonify(users)

@app.route('/api/get-validators')
def get_validators():
    if 'admin_email' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    validators = list(validators_collection.find({}, {'_id': 0, 'name': 1, 'email': 1,'categories': 1}))
    return jsonify(validators)

@app.route('/api/get-admins')
def get_admins():
    if 'admin_email' not in session:
        return jsonify({'error': 'Unauthorized'}), 401

    admins = list(admin_collection.find({}, {'_id': 0, 'admin_name': 1, 'admin_email': 1}))
    return jsonify(admins)


# Generate a password reset token
def generate_reset_token(email):
    s = URLSafeTimedSerializer(app.secret_key)
    return s.dumps(email, salt='password-reset-salt')


# Send the reset email
def send_reset_email(email):
    token = generate_reset_token(email)
    reset_link = url_for('reset_password', token=token, _external=True)
    print(f"Generated reset link: {reset_link}")  # Now this works!

    subject = "Password Reset Request"
    body = f"""To reset your password, click the following link:
{reset_link}

If you did not make this request, simply ignore this email.
"""

    msg = Message(subject=subject, recipients=[email], body=body)
    mail.send(msg)



def verify_reset_token(token, expiration=3600):
    s = URLSafeTimedSerializer(app.secret_key)
    try:
        email = s.loads(token, salt='password-reset-salt', max_age=expiration)
    except Exception:
        return None
    return email

# Helper function to validate file extensions
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Helper function to get default profile image
def get_default_profile_image():
    return url_for('static', filename='profile.webp')

@app.route('/send_email')
def send_email():
    try:
        msg = Message('Test Email', recipients=['recipient@example.com'])
        msg.body = 'This is a test email.'
        mail.send(msg)
        return "Email sent successfully!"
    except Exception as e:
        return f"Failed to send email: {e}"

@app.route('/oauth2callback')
def authorized():
    response = google.authorized_response()
    if response is None or response.get('access_token') is None:
        return 'Access denied: reason={} error={}'.format(
            request.args['reason'],
            request.args['error_reason']
        )

    session['google_token'] = (response['access_token'], '')
    user_info = google.get('userinfo')
    return 'Logged in as: ' + user_info.data['email']

@app.route('/authorize')
def authorize():
    token = google.authorize_access_token()
    resp = google.get('userinfo')
    user_info = resp.json()
    session['user'] = user_info
    return f"Hello, {user_info['email']}!"
@app.route('/refer')
def refer():
    if "email" in session:
        user = users_collection.find_one({"email": session["email"]})
        if user and user.get("image_id"):
            profile_img = url_for('get_profile_pic', image_id=user["image_id"])
        else:
            profile_img = get_default_profile_image()
    else:
        # Not logged in: show default image and possibly redirect or show login option
        user = None
        profile_img = get_default_profile_image()
        return redirect(url_for('login'))
    return render_template('Reference1.html', profile_img=profile_img, user=user)

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        form_type = request.form.get("form_type")

        if form_type == "login":
            email = request.form.get("email")
            password = request.form.get("password")
            user = users_collection.find_one({"email": email})

            if user and check_password_hash(user["password"], password):
                session["email"] = email
                session["name"] = user.get("name", "Unknown User")
                return redirect(url_for("work"))

            pending_user = pending_users_collection.find_one({"email": email})
            if pending_user:
                status = pending_user.get("status", "pending")
                if status == "rejected":
                    return alert_and_redirect("Admin has rejected your registration.", "login")
                else:
                    return alert_and_redirect("Kindly wait until the admin approves.", "login")

            return alert_and_redirect("Invalid credentials. Please try again.", "login")

        elif form_type == "register":
            name = request.form.get("name")
            email = request.form.get("email")
            password = request.form.get("password")

            existing_user = users_collection.find_one({"email": email})
            existing_pending = pending_users_collection.find_one({"email": email})

            if existing_user or existing_pending:
                return alert_and_redirect("You have already registered or are awaiting approval.", "login")

            hashed_password = generate_password_hash(password)
            pending_users_collection.insert_one({
                "name": name,
                "email": email,
                "password": hashed_password,
                "status": "pending"
            })

            return alert_and_redirect("Registration submitted. Please wait for admin approval.", "login")

    return render_template("login.html")

@app.route('/forgot_password', methods=['GET', 'POST'])
def forgot_password():
    if request.method == 'POST':
        email = request.form['email']
        user = users_collection.find_one({'email': email})  # or however you fetch the user
        if user:
            send_reset_email(user['email'])  # Only pass the email
            return render_template('forgot_password.html', message="A reset link has been sent to your email.")
        else:
            return render_template('forgot_password.html', message="No account found with that email.")
    return render_template('forgot_password.html')


@app.route('/reset_password/<token>', methods=['GET', 'POST'])
def reset_password(token):
    print(f"Received token: {token}")
    email = verify_reset_token(token)
    print(f"Decoded email: {email}")

    if not email:
        return 'The reset link is invalid or has expired.'

    if request.method == 'POST':
        new_password = request.form['password']
        hashed_password = generate_password_hash(new_password)
        users_collection.update_one({'email': email}, {'$set': {'password': hashed_password}})
        return redirect(url_for('login'))

    return render_template('reset_password.html')



@app.route('/work')
def work():
    if "email" in session:
        user = users_collection.find_one({"email": session["email"]})
        if user and user.get("image_id"):
            profile_img = url_for('get_profile_pic', image_id=user["image_id"])
        else:
            profile_img = get_default_profile_image()
    else:
        # Not logged in: show default image and possibly redirect or show login option
        user = None
        profile_img = get_default_profile_image()
        return redirect(url_for('login'))
    return render_template('work.html', name=user.get("name", "Unknown User"), profile_img=profile_img)

@app.route('/guide')
def guide():
    if "email" in session:
        user = users_collection.find_one({"email": session["email"]})
        if user and user.get("image_id"):
            profile_img = url_for('get_profile_pic', image_id=user["image_id"])
        else:
            profile_img = get_default_profile_image()
    else:
        # Not logged in: show default image and possibly redirect or show login option
        user = None
        profile_img = get_default_profile_image()
        return redirect(url_for('login'))
    return render_template('guide.html', profile_img=profile_img, user=user)

CORS(app)  # Allow frontend to call this API

responses = {
    "hi": ["Hello! How can I help you?", "Hi there! What do you need help with?", "Hey! What’s on your mind?"],
    "how are you": ["I'm just a bot, but I'm doing great! How about you?", "Feeling bot-tastic! What about you?", "I'm here and ready to assist you!"],
    "what is your name": ["I'm LuminaBot, your research assistant!", "Call me Lumina, your AI helper!"],
    "bye": ["Goodbye! Have a great day!", "See you soon!", "Take care!"],
    "thank you": ["You're welcome!", "Happy to help!", "Anytime!"],

}

import random

@app.route('/chat', methods=['POST'])
def chat():
    user_message = request.json.get("message", "").strip().lower()

    # Find a response or give a default one
    response_text = random.choice(responses.get(user_message, ["I'm not sure how to answer that. Can you rephrase?"]))

    return jsonify({"response": response_text})


GOOGLE_FORM_SCRIPT_URL = "https://script.google.com/macros/s/AKfycbzqngxHFiX4nTtqEdpc9p77_K0bHOGREt4n0X2BFSS-YJgoV-cV3wvhE_wDD-TzeCT-/exec"

# 📧 Email config (use your own credentials)
EMAIL_ADDRESS = "luminaassitant@gmail.com"
EMAIL_PASSWORD = "hfpe opim alde izqh"  # Use App Password if 2FA is on

@app.route('/upload', methods=['POST'])
def upload():
    try:
        data = request.get_json()
        questions = data.get('questions')
        user_email = data.get('user_email')

        if not questions or not user_email:
            return jsonify({'message': 'Missing questions or email'}), 400

        # ➡ Send to Google Script
        response = requests.post(GOOGLE_FORM_SCRIPT_URL, json={
            "questions": questions,
            "user_email": user_email
        })

        if response.status_code == 200:
            form_url = response.text.strip()

            # 📤 Send email with the form URL
            send_email(user_email, form_url)

            return jsonify({"form_url": form_url})
        else:
            return jsonify({
                "message": f"Google Script Error: {response.status_code} - {response.text}"
            }), 500

    except Exception as ex:
        return jsonify({'message': f"Unexpected error: {str(ex)}"}), 500

def send_email(to_email, form_url):
    msg = Message(subject='Your Google Form Questionnaire',
                  sender=app.config['MAIL_DEFAULT_SENDER'],
                  recipients=[to_email])
    msg.body = f"""
Hi,

Here is your questionnaire link:

{form_url}

You can open, edit, or share it as needed.

Regards,  
Team Lumina
"""
    try:
        mail.send(msg)
        print("Questionnaire email sent successfully.")
    except Exception as e:
        print("Failed to send questionnaire email:", e)

@app.route('/submit_evaluation', methods=['POST'])
def submit_evaluation():
    try:
        file = request.files['evaluated_file']
        user_email = request.form['email']
        category = request.form.get('category', 'General')

        if not file or not user_email:
            return jsonify({'message': 'Missing file or email'}), 400

        # Save file temporarily
        filename = secure_filename(file.filename)
        filepath = os.path.join('uploads', filename)
        os.makedirs('uploads', exist_ok=True)
        file.save(filepath)

        # ✉ Send evaluated file by email
        send_evaluated_email(user_email, filepath, filename, category)

        # Optional: remove file after sending
        os.remove(filepath)

        record_id = request.form.get('record_id')
        if record_id:
            assessment_collection.delete_one({"_id": ObjectId(record_id)})

        return jsonify({"message": "File evaluated and sent successfully."}), 200

    except Exception as ex:
        return jsonify({'message': f"Unexpected error: {str(ex)}"}), 500


def send_evaluated_email(to_email, file_path, filename, category):
    msg = EmailMessage()
    msg['Subject'] = f'Your Evaluated File - {category}'
    msg['From'] = EMAIL_ADDRESS
    msg['To'] = to_email
    msg.set_content(f"""
Hi,

Please find the evaluated file for the category "{category}" attached.

Best regards,  
Lumina Team
""")

    # Attach file
    with open(file_path, 'rb') as f:
        file_data = f.read()
        msg.add_attachment(file_data, maintype='application', subtype='octet-stream', filename=filename)

    # Send via Gmail
    with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
        smtp.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        smtp.send_message(msg)

import mimetypes

# Secure sender email and app password
EMAIL_ADDRESS_PLAGARISM = 'luminaassitant@gmail.com'
EMAIL_PASSWORD_PLAGARISM = 'hfpe opim alde izqh'  # App Password

@app.route('/upload_plagiarism', methods=['POST'])
def upload_plagiarism():
    try:
        user_email = session.get('email')
        if not user_email:
            return jsonify({'success': False, 'error': 'User email not found in session'}), 400

        research_file = request.files.get('researchDoc')
        payment_file = request.files.get('paymentReceipt')

        if not research_file or not payment_file:
            return jsonify({'success': False, 'error': 'Both files are required'}), 400

        # Compose the email
        msg = EmailMessage()
        msg['Subject'] = 'Plagiarism Submission'
        msg['From'] = EMAIL_ADDRESS_PLAGARISM  # ✅ Corrected here
        msg['To'] = EMAIL_ADDRESS_PLAGARISM     # send to yourself
        msg.set_content(f'User {user_email} has submitted files for plagiarism check.')

        # Attach research file
        research_data = research_file.read()
        research_type, _ = mimetypes.guess_type(research_file.filename)
        if not research_type:
            research_type = 'application/octet-stream'
        research_main, research_sub = research_type.split('/', 1)
        msg.add_attachment(research_data, maintype=research_main, subtype=research_sub, filename=research_file.filename)

        # Attach payment file
        payment_data = payment_file.read()
        payment_type, _ = mimetypes.guess_type(payment_file.filename)
        if not payment_type:
            payment_type = 'application/octet-stream'
        payment_main, payment_sub = payment_type.split('/', 1)
        msg.add_attachment(payment_data, maintype=payment_main, subtype=payment_sub, filename=payment_file.filename)

        # Send via Gmail SMTP
        with smtplib.SMTP_SSL('smtp.gmail.com', 465) as smtp:
            smtp.login(EMAIL_ADDRESS_PLAGARISM, EMAIL_PASSWORD_PLAGARISM)
            smtp.send_message(msg)

        return jsonify({'success': True}), 200

    except Exception as e:
        print(f"❌ Upload failed: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/export_word', methods=['POST'])
def export_word():
    data = request.get_json()
    questionnaire = data.get('questionnaire', [])

    doc = Document()
    doc.add_heading('Questionnaire', 0)

    for idx, item in enumerate(questionnaire, 1):
        doc.add_paragraph(f"{idx}. {item['question']}", style='List Number')
        for opt in item['options']:
            doc.add_paragraph(f"   - {opt}", style='List Bullet')

    # Save the doc to memory
    doc_stream = BytesIO()
    doc.save(doc_stream)
    doc_stream.seek(0)

    return send_file(doc_stream, as_attachment=True, download_name="questionnaire.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/logout")
def logout():
    session.pop("email", None)
    session.pop("name", None)  # Remove name from session on logout
    flash("Logged out successfully.", "info")
    return redirect(url_for("login"))

@app.route('/submit-query', methods=['POST'])
def submit_query():
    name = request.form['name']
    email = request.form['email']
    message = request.form['message']

    query_data = {
        'name': name,
        'email': email,
        'message': message
    }

    query_collection.insert_one(query_data)
    return alert_and_redirect("Your query submitted sucessfully ", "home")

@app.route('/get-queries')
def get_queries():
    queries = list(query_collection.find({}, {'_id': 0}))  # omit _id for display
    return jsonify(queries)

@app.route('/submit-feedback', methods=['POST'])
def submit_feedback():
    name = request.form['name']
    email = request.form['email']
    message = request.form['message']

    feedback_data = {
        'name': name,
        'email': email,
        'message': message
    }

    feedback_collection.insert_one(feedback_data)
    return alert_and_redirect("Your feedback submitted sucessfully ", "guide")

@app.route('/get-feedback')
def get_feedback():
    feedback = list(feedback_collection.find({}, {'_id': 0}))  # omit _id for display
    return jsonify(feedback)

@app.route('/stats', methods=['GET', 'POST'])
def pstats():
    global result_text, plot_stream_buffer
    result_html = None
    selected_tool = None
    image_file = None
    profile_img = url_for('static', filename='profile.webp')  # Default image
    error_message = None  # Declare this before the try block
    if "email" in session:
        user = users_collection.find_one({"email": session["email"]})
        if user and user.get("image_id"):
            profile_img = url_for('get_profile_pic', image_id=user["image_id"])
        else:
            profile_img = get_default_profile_image()
    else:
        # Not logged in: show default image and possibly redirect or show login option
        user = None
        profile_img = get_default_profile_image()
        return redirect(url_for('login'))

    if request.method == 'POST':
        selected_tool = request.form.get('analysis_tool')
        file = request.files.get('file')

        if not file or not selected_tool:
            return redirect(url_for('pstats'))

        df = pd.read_excel(file)
        result_text = ""  # Reset text result each time

        try:
            # Central Tendency
            if selected_tool in ['Mean Calculation', 'Median Calculation', 'Mode Calculation']:
                stat_func = {
                    'Mean Calculation': df.mean(numeric_only=True),
                    'Median Calculation': df.median(numeric_only=True),
                    'Mode Calculation': df.mode(numeric_only=True).iloc[0]
                }
                values = stat_func[selected_tool]
                result_html = "<table border='1' cellpadding='5'><tr><th>Column</th><th>Value</th></tr>"
                for col, val in values.items():
                    result_html += f"<tr><td>{col}</td><td>{val:.2f}</td></tr>"
                    result_text += f"{selected_tool.split()[0]} of {col}: {val:.2f}\n"
                result_html += "</table>"

            # Dispersion
            elif selected_tool in ['Variance', 'Standard Deviation']:
                if selected_tool == 'Variance':
                    values = df.var(numeric_only=True)
                else:
                    values = df.std(numeric_only=True)
                result_html = "<table border='1' cellpadding='5'><tr><th>Column</th><th>Value</th></tr>"
                for col, val in values.items():
                    result_html += f"<tr><td>{col}</td><td>{val:.2f}</td></tr>"
                    result_text += f"{selected_tool} of {col}: {val:.2f}\n"
                result_html += "</table>"

            # Hypothesis Testing (Multiple variables)
            elif selected_tool in ['T-Test', 'ANOVA', 'Chi-Square Test']:
                if selected_tool == 'T-Test':
                    # Perform T-Test for each column separately
                    result_html = "<h4>T-Test Results</h4><ul>"
                    for col in df.columns:
                        t_stat, p_val = stats.ttest_1samp(df[col].dropna(), 0)
                        result_html += f"<li><strong>{col}</strong>: T-Statistic: {t_stat:.4f}, P-Value: {p_val:.4f}</li>"
                        result_text += f"T-Test for {col}: T-Statistic: {t_stat:.4f}, P-Value: {p_val:.4f}\n"
                    result_html += "</ul>"

                elif selected_tool == 'ANOVA':
                    f_stat, p_val = stats.f_oneway(*[df[col] for col in df.columns if df[col].dtype != 'object'])
                    result_html = f"F-Statistic: {f_stat:.4f}, P-Value: {p_val:.4f}"
                    result_text = result_html

                elif selected_tool == 'Chi-Square Test':
                    # Perform Chi-Square for multiple columns
                    contingency_table = pd.crosstab(df.iloc[:, 0], df.iloc[:, 1])  # Using first two columns for contingency
                    chi2, p, dof, expected = stats.chi2_contingency(contingency_table)
                    result_html = f"""
                        <h4>Chi-Square Test Results</h4>
                        <p><strong>Chi2 Statistic:</strong> {chi2:.4f}</p>
                        <p><strong>p-value:</strong> {p:.4f}</p>
                        <p><strong>Degrees of Freedom:</strong> {dof}</p>
                        <p><strong>Expected Frequencies:</strong></p>
                        {pd.DataFrame(expected).to_html()}
                    """
                    result_text = f"Chi2: {chi2:.4f}\nP-Value: {p:.4f}\nDegrees of Freedom: {dof}\n{pd.DataFrame(expected).to_string()}"

            # Correlation & Regression
            elif selected_tool == 'Correlation Analysis':
                result_html = df.corr(numeric_only=True).to_html()
                result_text = df.corr(numeric_only=True).to_string()

            elif selected_tool == 'Regression Analysis':
                X = df.iloc[:, :-1]
                y = df.iloc[:, -1]
                X = sm.add_constant(X)
                model = sm.OLS(y, X).fit()
                result_html = model.summary().as_html()
                result_text = model.summary().as_text()

            # Visualizations
            elif selected_tool in ['Bar Chart', 'Pie Chart', 'Histogram', 'Scatter Plot', 'Line Graph']:
                plt.figure(figsize=(8, 6))
                if selected_tool == 'Bar Chart':
                    df.iloc[:, 0].value_counts().plot(kind='bar')
                elif selected_tool == 'Pie Chart':
                    plt.pie(df.iloc[:, 1], labels=df.iloc[:, 0], autopct='%1.1f%%')
                elif selected_tool == 'Histogram':
                    df.iloc[:, 0].plot(kind='hist')
                elif selected_tool == 'Scatter Plot':
                    plt.scatter(df.iloc[:, 0], df.iloc[:, 1])
                elif selected_tool == 'Line Graph':
                    df.iloc[:, 0].plot(kind='line')

                plt.title(selected_tool)
                plot_stream_buffer = BytesIO()
                plt.savefig(plot_stream_buffer, format='jpg')
                plot_stream_buffer.seek(0)
                plt.close()
                image_file = 'in-memory'

            # Factor Analysis
            elif selected_tool == 'Factor Analysis':
                numeric_df = df.select_dtypes(include=[np.number])
                fa = FactorAnalysis(n_components=2)
                factors = fa.fit_transform(numeric_df)
                factors_df = pd.DataFrame(factors, columns=['Factor 1', 'Factor 2'])
                result_html = "<h4>Factor Analysis (First 5 rows)</h4>" + factors_df.head().to_html()
                result_text = factors_df.head().to_string()


        except Exception as e:
            error_message = str(e)
            result_html = f"<p class='error'>Error: {error_message}</p>"
            result_text = result_html

    return render_template('stats.html', result=result_html, selected_tool=selected_tool,image_file=image_file, error_message=error_message, profile_img=profile_img)


@app.route('/plot_preview')
def plot_preview():
    global plot_stream_buffer
    if plot_stream_buffer:
        plot_stream_buffer.seek(0)
        return send_file(BytesIO(plot_stream_buffer.getvalue()), mimetype='image/jpeg')
    return "", 404

@app.route('/download_results')
def download_results():
    global result_text
    if result_text:
        stream = BytesIO()
        content = "Statistical Analysis Results\n" + "=" * 30 + "\n" + result_text
        stream.write(content.encode('utf-8'))
        stream.seek(0)
        return send_file(stream, as_attachment=True, download_name="results.txt", mimetype='text/plain')
    return "", 404

@app.route('/download_plot')
def download_plot():
    global plot_stream_buffer
    if plot_stream_buffer:
        plot_stream_buffer.seek(0)
        return send_file(BytesIO(plot_stream_buffer.getvalue()), mimetype='image/jpeg', as_attachment=True, download_name='plot.jpg')
    return "", 404

@app.route('/profile')
def profile():
    if "email" not in session:
        return redirect(url_for("login"))

    user = users_collection.find_one({'email': session['email']})  # ✅ Fixed line
    user_pdfs = list(pdfs_collection.find({"email": session["email"]}))

    for pdf in user_pdfs:
        pdf["file_id"] = str(pdf["file_id"])
        try:
            file_obj = fs.get(ObjectId(pdf["file_id"]))
            pdf["file_size"] = round(file_obj.length / 1024, 2)
            pdf["upload_date"] = file_obj.upload_date.strftime('%d-%m-%Y %H:%M')
        except Exception as e:
            pdf["file_size"] = "N/A"
            pdf["upload_date"] = "Unavailable"

    profile_img = get_default_profile_image()
    if user and user.get("image_id"):
        profile_img = url_for('get_profile_pic', image_id=user["image_id"])
    else:
        profile_img = get_default_profile_image()
    return render_template('profile.html', profile=user, profile_img=profile_img, pdfs=user_pdfs)

@app.route('/edit_profile', methods=['POST'])
def edit_profile():
    if 'email' not in session:
        return redirect(url_for('login'))

    name = request.form.get('name')
    experience = request.form.get('experience')
    profession = request.form.get('profession')

    users_collection.update_one(
        {'email': session['email']},
        {'$set': {
            'name': name,
            'experience': experience,
            'profession': profession
        }}
    )

    flash("Profile updated successfully!", "success")
    return redirect(url_for('profile'))

@app.route('/save_profile', methods=['POST'])
def save_profile():
    if "email" not in session:
        flash("Please log in to update your profile.", "danger")
        return redirect(url_for("login"))

    user = users_collection.find_one({"email": session["email"]})

    if not user:
        flash("User not found.", "danger")
        return redirect(url_for("login"))

    name = request.form.get('name')
    about = request.form.get('about')
    education = request.form.get('education')
    profession = request.form.get('profession')
    experience = request.form.get('experience')
    projects = request.form.get('projects')

    users_collection.update_one(
        {"email": session["email"]},
        {"$set": {
            "name": name,
            "about": about,
            "education": education,
            "profession": profession,
            "experience": experience,
            "projects": projects
        }}
    )

    if 'pdf_file' in request.files:
        pdf_files = request.files.getlist('pdf_file')
        for pdf_file in pdf_files:
            if pdf_file.filename:
                filename = secure_filename(pdf_file.filename)
                pdf_id = fs.put(pdf_file.read(), filename=filename, content_type=pdf_file.content_type)
                pdfs_collection.insert_one({
                    "email": session["email"],
                    "title": filename,
                    "file_id": pdf_id
                })

    flash("Profile updated successfully!", "success")
    return redirect(url_for("profile"))


@app.route('/upload_profile_pic', methods=['POST'])
def upload_profile_pic():
    if 'email' not in session:
        return redirect(url_for('login'))

    file = request.files.get('profile_pic')
    if file and file.filename != '':
        filename = secure_filename(file.filename)
        image_id = fs.put(file, filename=filename)

        users_collection.update_one({'email': session['email']}, {'$set': {'image_id': image_id}})
        flash("Profile picture updated!", "success")
    else:
        flash("No file selected.", "warning")

    return redirect(url_for('work'))



@app.route('/get_profile_pic/<image_id>')
def get_profile_pic(image_id):
    try:
        image_file = fs.get(ObjectId(image_id))
        return send_file(image_file, mimetype='image/jpeg')
    except:
        abort(404)

@app.route('/view_pdf/<pdf_id>')
def view_pdf(pdf_id):
    try:
        file = fs.get(ObjectId(pdf_id))
        return send_file(file, mimetype="application/pdf")
    except:
        return "PDF not found", 404

@app.route('/get_profile_picture/<image_id>')
def get_profile_picture(name):
    user = db.users_collection.find_one({"username": name})
    if user and user.get("profile_picture_id"):
        image = fs.get(user["profile_picture_id"])
        return Response(image.read(), mimetype='image/jpeg')
    else:
        return send_file("static/default_profile.png", mimetype='image/png')

@app.route('/download_pdf/<pdf_id>')
def download_pdf(pdf_id):
    try:
        file = fs.get(ObjectId(pdf_id))
        return send_file(file, as_attachment=True, download_name=file.filename)
    except:
        return "PDF not found", 404

@app.route('/delete_pdf/<pdf_id>')
def delete_pdf(pdf_id):
    if "email" not in session:
        return redirect(url_for("login"))
    try:
        pdf_id_obj = ObjectId(pdf_id)
        fs.delete(pdf_id_obj)
        pdfs_collection.delete_one({"file_id": pdf_id_obj, "email": session["email"]})
        flash("PDF deleted successfully!", "success")
        return redirect(url_for("profile"))
    except:
        flash("Error deleting PDF.", "danger")
        return redirect(url_for("profile"))

@app.route('/view-pdfs/<filename>')
def view_pdfs(filename):
    # Directory where PDFs are stored
    pdf_dir = os.path.join(app.root_path, 'static', 'thesis_pdfs')
    file_path = os.path.join(pdf_dir, filename)

    if not os.path.exists(file_path):
        return abort(404)

    return send_file(file_path, mimetype='application/pdf', as_attachment=False)

if __name__ == '__main__':
    app.run(debug=True)